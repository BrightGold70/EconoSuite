import pandas as pd
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_academic_table(document, csv_path):
    """
    Parses a CSV file and inserts a formatted academic table into the docx.
    """
    try:
        df = pd.read_csv(csv_path)
    except Exception as e:
        document.add_paragraph(f"[Error loading table: {csv_path}]")
        return
        
    rows = df.shape[0] + 1  # header
    cols = df.shape[1]
    
    table = document.add_table(rows=rows, cols=cols)
    table.style = 'Medium Shading 1 Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        
    for r_idx, row in df.iterrows():
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = str(value)

    document.add_paragraph()

def add_figure(document, img_path):
    """
    Inserts a figure (image file) into the docx.
    """
    if not os.path.exists(img_path):
        document.add_paragraph(f"[Error loading figure: File not found at {img_path}]")
        return
        
    try:
        # Add the picture and center it
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.0))
    except Exception as e:
        document.add_paragraph(f"[Error inserting figure: {str(e)}]")
