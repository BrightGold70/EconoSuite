import yaml
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
import docx.enum.text as enum_text
from .integrator import add_academic_table, add_figure

def build_document(base_dir):
    config_path = base_dir / "config.yaml"
    with open(config_path, "r") as f:
        config = yaml.safe_load(f)
        
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading(config.get("title", "Untitled Manuscript"), 0)
    title.alignment = enum_text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Authors
    authors = config.get("authors", [])
    if authors:
        author_text = ", ".join([a.get("name", "") for a in authors])
        p = doc.add_paragraph(author_text)
        p.alignment = enum_text.WD_ALIGN_PARAGRAPH.CENTER
        
    # Abstract
    doc.add_heading("Abstract", 1).alignment = enum_text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(config.get("abstract", ""))
    
    # JEL Codes
    jel = config.get("jel_codes", [])
    if jel:
        p = doc.add_paragraph()
        p.add_run("JEL Classification: ").bold = True
        p.add_run(", ".join(jel))
        
    doc.add_page_break()
    
    # Process markdown files sequentially
    md_files = sorted(list(base_dir.glob("*.md")))
    md_files = [m for m in md_files if m.name != "README.md"]
    
    table_pattern = re.compile(r"\{\{table:\s*(.+?)\}\}")
    figure_pattern = re.compile(r"\{\{figure:\s*(.+?)\}\}")
    
    for fpath in md_files:
        with open(fpath, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Headers
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            # Tables
            elif table_pattern.match(line):
                match = table_pattern.match(line)
                csv_rel_path = match.group(1).strip()
                csv_path = base_dir / csv_rel_path
                add_academic_table(doc, csv_path)
            # Figures 
            elif figure_pattern.match(line):
                match = figure_pattern.match(line)
                img_rel_path = match.group(1).strip()
                img_path = base_dir / img_rel_path
                add_figure(doc, img_path)
            # Paragraphs
            else:
                doc.add_paragraph(line)
                
    output_filename = f"{config.get('title', 'manuscript').replace(' ', '_')}.docx"
    output_path = base_dir / output_filename
    doc.save(str(output_path))
    return str(output_path)
